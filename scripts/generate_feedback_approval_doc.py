from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "articles-managers-feedback-approval-pack.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_heading_run(paragraph, text, size, color=None, bold=True):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def format_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size in [("Title", 22), ("Heading 1", 16), ("Heading 2", 12.5), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)


def add_title_block(doc, prepared_on):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading_run(p, "Articles and Managers Feedback", 22, "1F2937")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading_run(p2, "Approval Pack for Sharing", 13, "4B5563", bold=False)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run(f"Prepared on {prepared_on.strftime('%d %B %Y')}")
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string("6B7280")


def add_meta_table(doc, prepared_on):
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    rows = [
        ("Document purpose", "Provide an approval-ready overview of the two feedback modules used in the Timesheet app."),
        ("Feedback tracks", "Articles Feedback and Managers Feedback."),
        ("Audience", "Managers, partners, and articled assistants."),
        ("Document date", prepared_on.strftime("%d %B %Y")),
    ]
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        set_cell_shading(table.cell(i, 0), "D9EAF7")
        for cell in (table.cell(i, 0), table.cell(i, 1)):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = "Aptos"
        table.cell(i, 0).paragraphs[0].runs[0].bold = True


def add_section_overview(doc):
    doc.add_heading("1. Purpose and scope", level=1)
    doc.add_paragraph(
        "This pack is intended for review and approval before sharing the feedback forms with users. "
        "It keeps the wording aligned to the application language and separates the two anonymous feedback flows."
    )
    add_bullet(doc, "Articles Feedback is completed by managers and partners about articled assistants.")
    add_bullet(doc, "Managers Feedback is completed by articled assistants about managers, partners, workload, and culture.")
    add_bullet(doc, "Both forms are anonymous at the report level and should be shared only with authorized reviewers.")


def add_articles_feedback(doc):
    doc.add_heading("2. Articles Feedback", level=1)
    doc.add_paragraph(
        "Audience: managers and partners reviewing articled assistants.\n"
        "Goal: capture professional performance, growth, team behavior, and overall recommendation."
    )

    sections = [
        (
            "Professionalism and delivery",
            [
                "How professional and responsive is the articled assistant in daily work?",
                "How often are deadlines missed without proactive communication?",
                "What strengths do you most often observe?",
                "What recurring issue most affects delivery quality?",
            ],
        ),
        (
            "Learning and growth",
            [
                "How receptive is the assistant to feedback and correction?",
                "How would you rate their learning speed?",
                "Which areas need the most support?",
                "What should be done to accelerate their growth?",
            ],
        ),
        (
            "Team behaviour",
            [
                "How well do they work with seniors and peers?",
                "How responsive are they when the team needs a quick turnaround?",
                "What team traits do they show consistently?",
                "What one concern would you want leadership to know?",
            ],
        ),
        (
            "Overall",
            [
                "How likely are you to recommend reassigning work to this assistant?",
                "Overall, how would you rate this articled assistant right now?",
                "What should leadership know that is not obvious from daily work?",
            ],
        ),
    ]

    for title, items in sections:
        doc.add_heading(title, level=2)
        for item in items:
            add_bullet(doc, item)

    doc.add_paragraph(
        "Approval note: This module should remain concise, respectful, and suitable for management review."
    )


def add_managers_feedback(doc):
    doc.add_heading("3. Managers Feedback", level=1)
    doc.add_paragraph(
        "Audience: articled assistants providing anonymous feedback on managers, partners, and firm culture.\n"
        "Goal: surface workload, guidance, behavior, culture, and overall experience."
    )

    sections = [
        (
            "Workload and hours",
            [
                "How fairly is work distributed among articled assistants?",
                "How often are you required to work beyond standard hours without prior notice?",
                "During peak season, what support do you receive?",
                "Are you given adequate time to complete assignments with quality?",
                "Describe a situation where workload felt unmanageable. What would have helped?",
            ],
        ),
        (
            "Guidance and learning",
            [
                "How well do seniors explain the purpose and context of tasks they assign?",
                "How often does the firm conduct internal training sessions?",
                "How accessible and approachable are managers when guidance is needed?",
                "Is there a specific topic you needed guidance on but could not find support for?",
            ],
        ),
        (
            "Management behaviour",
            [
                "How respectfully do partners and seniors communicate with you?",
                "Which behaviors have you personally experienced or witnessed?",
                "When you ask a question or raise a concern, how does your supervisor respond?",
                "How fair and transparent is the performance feedback you receive?",
                "Describe a management behavior that has negatively affected motivation or wellbeing.",
            ],
        ),
        (
            "Culture and environment",
            [
                "How safe do you feel raising concerns or disagreeing with a senior?",
                "Which words best describe the current culture of the firm?",
                "How often do seniors acknowledge your contribution or good work?",
                "How valued do you feel as a future professional, not just as labour?",
                "What one cultural change would make you recommend this firm to the next batch?",
            ],
        ),
        (
            "Overall experience",
            [
                "How likely are you to refer a friend to do their articleship here?",
                "Are you considering leaving before completing your articleship?",
                "If you stay, what is the primary reason?",
                "Overall, how would you rate your articleship experience so far?",
                "What should the firm's leadership know that they probably do not?",
            ],
        ),
    ]

    for title, items in sections:
        doc.add_heading(title, level=2)
        for item in items:
            add_bullet(doc, item)

    doc.add_paragraph(
        "Approval note: This module is intentionally anonymous and should be presented as feedback for improvement, not attribution."
    )


def add_approval_table(doc):
    doc.add_heading("4. Approval", level=1)
    doc.add_paragraph("Use the table below for review and sign-off before sharing externally.")

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Reviewer", "Role", "Decision / Remarks"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_shading(cell, "1F2937")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
            run.font.size = Pt(10.5)

    set_repeat_table_header(table.rows[0])

    rows = [
        ("Prepared by", "Operations / HR", "Draft complete"),
        ("Reviewed by", "Manager / Partner", "____________________________"),
        ("Approved by", "Final approver", "____________________________"),
    ]
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10.5)


def add_footer_note(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared for internal review and approval.")
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string("6B7280")


def build_doc():
    doc = Document()
    format_document(doc)

    prepared_on = date.today()
    add_title_block(doc, prepared_on)
    doc.add_paragraph("")
    add_meta_table(doc, prepared_on)
    doc.add_paragraph("")
    add_section_overview(doc)
    add_articles_feedback(doc)
    add_managers_feedback(doc)
    add_approval_table(doc)
    add_footer_note(doc)
    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
